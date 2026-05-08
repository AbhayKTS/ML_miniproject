from __future__ import annotations

from pathlib import Path
import re
from typing import List, Dict

from docx import Document

TEMPLATE_PATH = Path(r"t:\ml_project\final_report_format.docx")
OUTPUT_PATH = Path(r"t:\ml_project\Chhaya_Project_Report.docx")

PROJECT_TITLE = "Enhancing Creative Content Generation Using Adaptive Generative AI (Chhaya AI)"
TEAM = ["Ansh Yadav", "Abhay Kumar", "Adarsh Katara"]
GUIDE_NAME = "Project Guide"
GROUP_NO = "01"

SECTION_CONTENT: Dict[str, str] = {
    "DECLARATION": (
        "We, {team}, students of B.Tech in Computer Science and Engineering (Artificial Intelligence and Machine Learning) at "
        "GLA University, Mathura, hereby solemnly declare that the project report entitled \"{title}\" submitted in partial "
        "fulfillment of the requirements for the award of the degree of Bachelor of Technology is an authentic record of our "
        "original work carried out during the eighth semester under the valuable guidance and supervision of {guide}, Department "
        "of Computer Engineering & Applications.\n\n"
        "We further declare that the work presented in this report is a result of our own efforts, research, and contributions, "
        "except where explicit references have been made. The findings, designs, implementations, and conclusions drawn are based "
        "on the project conducted by us. This work has not been submitted elsewhere, either in part or in full, for the award of any "
        "degree, diploma, or academic qualification.\n\n"
        "We also acknowledge that we have adhered to the ethical guidelines, academic integrity, and project norms set by the "
        "University throughout the course of this project."
    ),
    "CERTIFICATE": (
        "This is to certify that the project report entitled \"{title}\" submitted by {team} in partial fulfillment of the "
        "requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering (Artificial "
        "Intelligence and Machine Learning) from GLA University, Mathura, is a bona fide record of the original work carried out by "
        "them under my supervision and guidance during the academic year 2025-2026.\n\n"
        "The work embodied in this report is genuine to the best of my knowledge and has not been submitted, either in part or in "
        "full, for the award of any other degree or diploma in this or any other institution. The students have fulfilled all the "
        "requirements as per the academic guidelines prescribed by the University.\n\n"
        "{guide}\n"
        "Department of Computer Engineering & Applications\n"
        "GLA University, Mathura"
    ),
    "ACKNOWLEDGEMENT": (
        "The successful completion of any project is never the work of a single individual. It is the culmination of collective "
        "effort, guidance, and support from many quarters. We take this opportunity to express our profound gratitude to all those "
        "who have contributed directly or indirectly to the successful execution of this project.\n\n"
        "We extend our sincere thanks to our project guide for continuous guidance, constructive criticism, and encouragement. We "
        "are grateful to the Head of Department and all faculty members for providing the necessary infrastructure and academic "
        "environment. Finally, we thank our parents and friends for their unwavering moral support and motivation."
    ),
    "ABSTRACT": (
        "Chhaya AI is an adaptive multi-modal creative intelligence system that assists users in generating text, image, and audio "
        "content through a unified workflow. The platform blends prompt engineering with a persistent creative memory and feedback "
        "loops to keep outputs aligned with user intent, cultural context, and stylistic preferences.\n\n"
        "The system integrates specialized engines through a Hugging Face API key and provider routing (MistralAI for text generation, "
        "Stability AI for image generation, and ESPnet for audio synthesis). A Node.js API orchestrates generation requests, applies "
        "constraints, and records adaptive profiles, while a React-based UI provides workspaces for iterative refinement.\n\n"
        "This report presents the problem motivation, architecture, methodology, and evaluation of Chhaya AI, demonstrating how a "
        "modular, feedback-driven pipeline can enhance creative production quality, consistency, and speed."
    ),
    "Area of Project": (
        "This project falls under the following research and engineering domains: Generative Artificial Intelligence (using large "
        "language models and diffusion models for creative content synthesis), Adaptive Machine Learning (building systems that evolve "
        "from user feedback and behavioral signals), Full-Stack Web Engineering (designing a complete client-server application), "
        "Human-Computer Interaction (creating responsive interfaces for AI-powered creativity), and Multi-Modal AI Systems "
        "(integrating text, image, and audio pipelines into one system)."
    ),
    "Technical Keywords": (
        "Generative AI, Adaptive Learning, Multimodal Content Generation, Creative Memory, Hugging Face API Key, MistralAI, Stability AI, "
        "ESPnet, Node.js, Express.js, React, TypeScript, Supabase, PostgreSQL, Firebase, JWT Authentication, CSRF Protection, REST API, "
        "Cross-Modal Planning, Feedback Loop, Personalization Engine"
    ),
    "Project Idea": (
        "The central idea of Chhaya AI is to create an intelligent adaptive creative studio where the AI evolves with each user's "
        "creative style. Instead of static output generation, Chhaya AI stores a Creative Memory Profile for every user containing "
        "preferred tone, recurring themes, visual style, audio aesthetics, and cultural context.\n\n"
        "On every new generation, the system uses this profile to personalize output. When users provide ratings or edits, the Adaptive "
        "Learning Engine updates the profile through model-assisted interpretation or rule-based heuristics. This creates a continuous "
        "feedback loop that improves alignment over time.\n\n"
        "The current implementation uses a Hugging Face API key as the unified access layer with provider-level routing: MistralAI for "
        "text generation, Stability AI for image generation, and ESPnet for audio generation."
    ),
    "Motivation of the Project": (
        "Creative teams often juggle multiple tools for scripting, visuals, and audio. This fragmentation leads to context loss and "
        "inconsistent style. Chhaya AI addresses this by providing a unified pipeline where intent, tone, and constraints persist "
        "across modalities, improving coherence and reducing iteration time."
    ),
    "Literature Survey": (
        "Recent advances in large language models and diffusion-based image models have enabled rapid creative generation. However, "
        "most systems remain single-modal and lack long-term memory. Research in prompt engineering, retrieval-augmented generation, "
        "and preference learning highlights the value of structured controls and feedback loops. Chhaya AI incorporates these ideas "
        "by combining a reasoning layer, creative memory, and modality-specific engines."
    ),
    "Problem Statement": "Enhancing Creative Content Generation Using Adaptive Generative AI (Chhaya AI).",
    "Goals and objectives": (
        "Develop a multi-modal creative platform, maintain persistent creative memory, enforce cultural and tonal constraints, and "
        "enable feedback-driven refinement. Ensure that the platform is usable, scalable, and produces consistent outputs across "
        "text, image, audio, and video."
    ),
    "Statement of scope": (
        "The scope includes a web-based UI, API orchestration, prompt generation, and integration with third-party generation "
        "services. Non-goals include training proprietary foundation models or building production-grade deployment pipelines."
    ),
    "Major Constraints": (
        "External API rate limits, model costs, and latency are primary constraints. Additional constraints include safe content "
        "filtering, storage limits for generated assets, and maintaining prompt privacy."
    ),
    "Methodologies of Problem solving": (
        "We used an iterative development methodology with rapid prototyping and continuous integration of user feedback. The "
        "system was designed using modular services to allow independent improvements in each engine."
    ),
    "Analysis": (
        "Requirement analysis focused on user flows for prompt creation, generation, and feedback. We identified critical system "
        "components: authentication, generation orchestration, creative memory, and analytics."
    ),
    "Design": (
        "The design uses a layered architecture: UI, API gateway, adaptive processing, and data services. Each layer exposes clear "
        "interfaces, enabling substitution of model providers or storage systems."
    ),
    "Development": (
        "Development was done in Node.js for the backend and React for the frontend. We integrated a Hugging Face API key-based "
        "provider routing setup with MistralAI, Stability AI, and ESPnet through service adapters and created a shared schema for "
        "generation requests."
    ),
    "Evaluation": (
        "Evaluation considered prompt quality, cross-modal consistency, responsiveness, and usability. Tests focused on API "
        "responses, memory updates, and correct routing to each generation engine."
    ),
    "Outcome": (
        "A working prototype that generates multi-modal creative outputs, adapts to feedback, and provides a cohesive workflow in a "
        "single interface."
    ),
    "Applications": (
        "Content creators, marketing teams, social media managers, and educators can use Chhaya AI to generate scripts, visuals, and "
        "audio assets with consistent style and context."
    ),
    "Hardware Resources Required": (
        "Standard development laptop (8GB+ RAM), internet access for API calls, and optional GPU for local experimentation."
    ),
    "Software Resources Required": (
        "Node.js, npm, React (Vite), Firebase Auth, Hugging Face API key access, and model providers (MistralAI, Stability AI, ESPnet)."
    ),
    "Project Model Analysis": (
        "A hybrid Agile model was adopted. Iterations were structured around delivering functional workspaces (text, image, and audio) "
        "and integrating user feedback into creative memory."
    ),
    "Reconciled Estimates": (
        "Effort estimates considered API integration time, UI development, and testing. Revisions were applied after early spikes to "
        "balance integration complexity with UI polish."
    ),
    "Cost Estimation using COCOMO(Constructive Cost) Model": (
        "Using a small-scale organic model, the estimated effort is modest due to heavy reliance on managed AI services. The main "
        "cost drivers are integration and UI/UX iteration rather than algorithmic implementation."
    ),
    "Risk Management w.r.t. NP Hard analysis": (
        "Key risks include dependency on external AI services, cost fluctuations, and content safety. Mitigations include caching, "
        "rate limiting, fallback behavior, and modular service design."
    ),
    "Risk Identification": "Identify provider outages, API key misconfiguration, and inconsistent output quality.",
    "Risk Analysis": "Prioritize risks by likelihood and impact, emphasizing service availability and compliance constraints.",
    "Project Schedule": (
        "The project schedule followed phases: requirements analysis, architecture design, prompt workflow implementation, API "
        "integration, and validation with sample outputs."
    ),
    "Project task set": (
        "Task sets included UI design, backend service integration, authentication, creative memory, and generation testing."
    ),
    "Timeline chart": "A Gantt-style timeline was maintained with weekly milestones for each module and integration checkpoint.",
    "Introduction": (
        "This document describes the design and implementation of Chhaya AI, an adaptive multi-modal creative platform. The system "
        "coordinates multiple AI services while maintaining stylistic continuity and user feedback loops."
    ),
    "Purpose and Scope of Document": (
        "This report documents the system requirements, architecture, design decisions, implementation details, testing, and "
        "deployment considerations for Chhaya AI."
    ),
    "Use Case View": (
        "Primary use cases include: create prompt, generate text, generate image prompt, generate audio prompt, "
        "submit feedback, and update creative memory."
    ),
    "Functional Model and Description": (
        "The functional model maps user requests to generation services. Requests are validated, enriched with memory and controls, "
        "routed to the appropriate engine, and stored for feedback learning."
    ),
    "Data Flow Diagram": (
        "Data flows from the UI to the API gateway, into the creative engine for intent analysis, then to modality-specific services, "
        "and finally to storage and feedback modules."
    ),
    "Activity Diagram:": (
        "User selects workspace -> enters prompt -> sets controls -> triggers generation -> views output -> submits feedback -> "
        "memory updates."
    ),
    "Non Functional Requirements:": (
        "The system must be responsive, secure, and reliable. Target latency for prompt generation is under a few seconds. Logs and "
        "monitoring should provide traceability and error diagnosis."
    ),
    "Sequence Diagram": (
        "UI sends request -> API validates -> memory service enriches -> generation engine produces output -> API stores result -> UI "
        "displays output."
    ),
    "System Architecture": (
        "Chhaya AI uses a layered architecture with UI, API gateway, adaptive processing, and data services. Each generation engine "
        "is encapsulated as a service adapter to allow future replacement without large refactors."
    ),
    "Architectural Design": (
        "Modules include authentication, prompt orchestration, creative memory, feedback analytics, and external AI connectors."
    ),
    "Module 1 : Data-set Gathering": (
        "Module 1 focuses on user input capture and prompt intake from the UI, including creative briefs, tone controls, and "
        "constraints."
    ),
    "Module 2 : Pre-processing": (
        "Module 2 normalizes user prompts, merges creative memory signals, and applies constraint alignment before generation."
    ),
    "Module 3: Data-set split": (
        "Module 3 routes requests to the correct modality-specific engine (text, image, audio, video) based on the workflow stage."
    ),
    "Module 4: Model Architecture": (
        "Module 4 encapsulates the multi-engine routing layer via Hugging Face key-based access: MistralAI for text, Stability AI for "
        "images, and ESPnet for audio."
    ),
    "Module 5: Hyper-parameter tuning": (
        "Module 5 manages adaptive controls (originality, complexity, cultural context) and tunes prompts according to feedback."
    ),
    "Tools and Technologies Used": (
        "React (Vite), Node.js (Express), Firebase Auth, Hugging Face API key integrations, MistralAI, Stability AI, ESPnet, Supabase "
        "(optional), and Postgres-compatible storage."
    ),
    "Planning": "Agile sprints with weekly integration checkpoints and UI reviews.",
    "UML Tools": "Draw.io and Figma for diagrams and UI planning.",
    "Programming Languages": "JavaScript, TypeScript, and Python for tooling scripts.",
    "Programming Frameworks": "React, Express.js, and Vite.",
    "IDE": "Visual Studio Code.",
    "Versioning Control": "Git and GitHub.",
    "Cloud Services": "Firebase Auth, Hugging Face-hosted provider access, and optional object storage.",
    "Application and web servers:": "Node.js runtime with Express API and Vite dev server.",
    "Libraries": "firebase, zod, axios/fetch, Hugging Face client utilities, and project utility libraries.",
    "Algorithm Details": (
        "The adaptive reasoning workflow combines intent analysis, constraint alignment, cultural context checks, and feedback-driven "
        "memory updates."
    ),
    "Dataset Details": (
        "No offline dataset training is required; the system relies on user prompts and feedback as dynamic data inputs."
    ),
    "Preprocessing Details": (
        "Prompts are normalized, control parameters validated, and memory signals merged to ensure consistent formatting and safety."
    ),
    "Model Details": (
        "Text generation uses MistralAI, image generation uses Stability AI, and audio generation uses ESPnet, all accessed via a "
        "Hugging Face API key workflow."
    ),
    "Model Training Details": (
        "Chhaya AI does not train custom models; it orchestrates external engines and updates user profiles using stored feedback."
    ),
    "Model Prediction Details": (
        "Prediction is performed via API calls to the chosen engine, with results stored and tagged by modality and context."
    ),
    "Type of Testing Used": (
        "Unit tests for services, API integration tests for endpoints, UI smoke tests, and manual workflow validation."
    ),
    "Test Cases and Test Results": (
        "Test cases include authentication, generation success/failure handling, memory update correctness, and UI state updates. "
        "Results show stable responses across common workflows."
    ),
    "Screen shots": (
        "Screenshots include the Landing page, Dashboard, Prompt Studio (text/image/audio), Media Generation page, and Feedback "
        "panel outputs."
    ),
    "Deployment": (
        "The system supports local development using Node.js and Vite. Production deployment would use containerized services and "
        "secure environment variables for API keys."
    ),
    "Conclusion": (
        "Chhaya AI demonstrates how adaptive creative intelligence can unify multi-modal generation into a coherent workflow. The "
        "system improves consistency, reduces iteration time, and enables personalized creative output."
    ),
    "Future Scope": (
        "Future work includes improved memory embeddings, real-time collaboration, richer analytics, and on-device inference options."
    ),
}

HEADING_RENAMES = {
    "Deepfake Video Detection And Analysis": PROJECT_TITLE,
}

HEADING_OVERRIDES = {
    "Module 1 : Data-set Gathering": "Module 1: Prompt Intake",
    "Module 2 : Pre-processing": "Module 2: Intent & Constraint Alignment",
    "Module 3: Data-set split": "Module 3: Engine Routing",
    "Module 4: Model Architecture": "Module 4: Multi-Engine Generation",
    "Module 5: Hyper-parameter tuning": "Module 5: Adaptive Controls",
}


def split_into_paragraphs(text: str, count: int) -> List[str]:
    parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not parts:
        parts = [text.strip()]

    if count <= 1:
        return [" ".join(parts)]

    sentences = []
    for part in parts:
        for sent in re.split(r"(?<=[.!?])\s+", part):
            if sent.strip():
                sentences.append(sent.strip())

    if not sentences:
        sentences = parts

    bucket_count = max(1, count)
    buckets = [""] * bucket_count
    for i, sent in enumerate(sentences):
        idx = i % bucket_count
        if buckets[idx]:
            buckets[idx] += " " + sent
        else:
            buckets[idx] = sent

    buckets = [b if b.strip() else "" for b in buckets]
    return buckets


def replace_text(paragraph_text: str) -> str:
    text = paragraph_text
    for old, new in HEADING_RENAMES.items():
        text = text.replace(old, new)
    text = text.replace("Group No.: 02", f"Group No.: {GROUP_NO}")
    text = re.sub(r"Priyanshu Bhardwaj\(.*?\)", TEAM[0], text)
    text = re.sub(r"Rahul Gangwar\(.*?\)", TEAM[1], text)
    text = re.sub(r"Sachin Kumar\(.*?\)", TEAM[2], text)
    text = text.replace("Mr. Preshit Mangesh Desai", GUIDE_NAME)
    return text


def main() -> None:
    doc = Document(TEMPLATE_PATH)

    # Replace static text occurrences.
    for p in doc.paragraphs:
        if p.text:
            p.text = replace_text(p.text)

    # Update cover title directly if needed.
    for p in doc.paragraphs:
        if p.text.strip() == "Deepfake Video Detection And Analysis":
            p.text = PROJECT_TITLE

    # Heading renames
    for p in doc.paragraphs:
        if p.text.strip() in HEADING_OVERRIDES:
            p.text = HEADING_OVERRIDES[p.text.strip()]

    # Build section ranges using headings and special markers.
    markers = {"ACKNOWLEDGEMENT": None, "ABSTRACT": None}
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() in markers:
            markers[p.text.strip()] = i

    headings = []
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip():
            headings.append((i, p.text.strip()))

    for key, idx in markers.items():
        if idx is not None:
            headings.append((idx, key))

    headings.sort(key=lambda x: x[0])

    # Replace content between headings.
    for idx, (start, heading_text) in enumerate(headings):
        if heading_text not in SECTION_CONTENT:
            continue
        end = headings[idx + 1][0] if idx + 1 < len(headings) else len(doc.paragraphs)
        content = SECTION_CONTENT[heading_text].format(
            team=", ".join(TEAM),
            title=PROJECT_TITLE,
            guide=GUIDE_NAME
        )
        paras = split_into_paragraphs(content, max(1, end - start - 1))
        for j, para_text in enumerate(paras):
            target_idx = start + 1 + j
            if target_idx >= end:
                break
            doc.paragraphs[target_idx].text = para_text
        # Clear any remaining paragraphs in the section.
        for k in range(start + 1 + len(paras), end):
            doc.paragraphs[k].text = ""

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
